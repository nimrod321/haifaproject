from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cv():
    doc = Document()

    # Title
    name = doc.add_heading('Nimrod Allouche', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph('Email: nimrod.allouche@gmail.com | Phone: 0587885760')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        "Third-year Computer Engineering student at the Technion (GPA: 80.2) with a strong foundation in both software development and electrical engineering. Passionate about system architecture, low-level programming, hardware verification, and digital logic design. Eager to apply expertise in systems programming, hardware description languages, and algorithm development to a challenging engineering role."
    )

    # Education
    doc.add_heading('Education', level=1)
    p = doc.add_paragraph()
    p.add_run('Technion - Israel Institute of Technology').bold = True
    p.add_run(' – Haifa, Israel (2023 – Present)\n')
    p.add_run('B.Sc, Computer Engineering\n')
    p.add_run('GPA: ').bold = True
    p.add_run('80.2/100\n')
    p.add_run('Honors: ').bold = True
    p.add_run("Dean's List (August 2024)\n")
    p.add_run('Relevant Coursework: ').bold = True
    p.add_run('Data Structures, Introduction to Systems Programming, Digital Systems, Introduction to C, Hardware Description Languages.')

    # Technical Skills
    doc.add_heading('Technical Skills', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('Programming & HDLs: ').bold = True
    p2.add_run('JavaScript, HTML, CSS, C, C++, Python, Java, Assembly, Verilog, SystemVerilog\n')
    p2.add_run('Frameworks & Technologies: ').bold = True
    p2.add_run('WebSockets (Socket.IO), Flask\n')
    p2.add_run('Developer Tools: ').bold = True
    p2.add_run('Git, Linux, VS Code, SQL\n')
    p2.add_run('Concepts & Electrical Knowledge: ').bold = True
    p2.add_run('Algorithms, Data Structures, Object-Oriented Programming (OOP), Finite State Machines (FSM), Hardware Verification, Clock Domain Crossing (CDC), Low-Level System Architecture.')

    # Experience
    doc.add_heading('Experience', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('Full-Stack Developer (Research Assistant) | Haifa University ').bold = True
    p3.add_run('(April 2024 – Present)')
    doc.add_paragraph("Engineered a real-time multiplayer web application from scratch using Python (Flask) and JavaScript, supporting complex game theory experiments (Prisoner's Dilemma) for university research.", style='List Bullet')
    doc.add_paragraph("Architected the complete backend infrastructure, implementing bi-directional real-time communication via WebSockets (Socket.IO) to synchronize concurrent player interactions and AI bot logic.", style='List Bullet')
    doc.add_paragraph("Designed a custom 2D procedural rendering engine using the HTML5 Canvas API, including collision detection, dynamic tilemap generation, and interactive UI overlays.", style='List Bullet')
    doc.add_paragraph("Developed robust database schemas using SQLite/SQL to track user authentication, virtual economies, and game state logs, automating the export of analytics directly to Excel.", style='List Bullet')

    p4 = doc.add_paragraph()
    p4.add_run('Electronic Warfare Combatant ').bold = True
    p4.add_run('(Apr 2019 – Dec 2021) Israel Defense Forces')
    doc.add_paragraph("Operated, troubleshot, and maintained advanced electrical and electronic warfare (EW) systems in a high-pressure field environment.", style='List Bullet')
    doc.add_paragraph("Analyzed signal intelligence (SIGINT) to identify and report on operational threats.", style='List Bullet')
    doc.add_paragraph("Collaborated with a team to ensure 24/7 operational readiness and electrical system integrity.", style='List Bullet')

    # Projects
    doc.add_heading('Projects', level=1)
    p5 = doc.add_paragraph()
    p5.add_run('Linux Kernel & System Architecture | Technion').bold = True
    doc.add_paragraph("Modified and compiled the Linux Kernel (vX.X), implementing custom system calls in C to execute a comprehensive Linux kernel modification and edit kernel headers within a virtual machine environment.", style='List Bullet')

    p6 = doc.add_paragraph()
    p6.add_run('Digital Logic & Hardware Design').bold = True
    doc.add_paragraph("Designed and simulated complex finite state machines (FSMs) using hardware description languages, including implementing custom modulo logic circuitry.", style='List Bullet')

    # Honors & Awards
    doc.add_heading('Honors & Awards', level=1)
    p7 = doc.add_paragraph()
    p7.add_run("Technion Dean's List (August 2024): ").bold = True
    p7.add_run("Awarded for outstanding academic achievement during the second semester.")

    # Languages
    doc.add_heading('Languages', level=1)
    doc.add_paragraph("Hebrew (Native)", style='List Bullet')
    doc.add_paragraph("English (Professional Working Proficiency)", style='List Bullet')

    # Hobbies
    doc.add_heading('Hobbies', level=1)
    doc.add_paragraph("Chess, solving riddles, high-intensity fitness.", style='List Bullet')

    doc.save('Nimrod_Allouche_CV.docx')
    print("CV generated successfully: Nimrod_Allouche_CV.docx")

if __name__ == '__main__':
    create_cv()
